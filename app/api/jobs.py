"""Job status & result endpoints."""

import io
import textwrap

from docx import Document
from fastapi.responses import PlainTextResponse
from fastapi import APIRouter, Depends, HTTPException, Response
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from rq.exceptions import NoSuchJobError
from rq.job import Job as RQJob
from sqlalchemy.orm import Session

from app.api.rate_limit import rate_limit
from app.config import settings
from app.db.database import get_db
from app.db.models import Job, User
from app.api.deps import get_current_user
from app.queue import enqueue_task, get_queue

router = APIRouter(prefix="/v1/jobs", tags=["jobs"])
retry_limiter = rate_limit(settings.job_submit_rate_limit_per_minute, 60, "retry")


def _render_sections(job: Job, template: str = "default") -> list[tuple[str, list[str]]]:
    source = job.source_meta or {}
    transcript = job.transcript or {}
    summary = job.summary or {}
    source_lines = [f"Job ID: {job.id}", f"Source Type: {job.source_type}"]
    if source.get("title"):
        source_lines.append(f"Title: {source.get('title')}")
    if source.get("filename"):
        source_lines.append(f"File: {source.get('filename')}")
    if source.get("url"):
        source_lines.append(f"URL: {source.get('url')}")

    key_points = [str(p) for p in (summary.get("key_points") or [])] or ["—"]
    action_items = [str(a) for a in (summary.get("action_items") or [])] or ["—"]
    timestamps = [
        f"{ts.get('t', '--:--:--')} {ts.get('label', '')}".strip()
        for ts in (summary.get("timestamps") or [])
    ] or ["—"]

    outline_lines = []
    outline = summary.get("outline") or []
    if outline:
        for sec in outline:
            title = sec.get("title") if isinstance(sec, dict) else str(sec)
            outline_lines.append(f"{title or 'Section'}")
            points = sec.get("points", []) if isinstance(sec, dict) else []
            for point in points:
                outline_lines.append(f"  - {point}")
    else:
        outline_lines = ["—"]

    tl_dr = summary.get("tl_dr") or "—"
    transcript_text = transcript.get("text") or ""

    if template == "meeting_notes":
        return [
            ("Meeting Overview", source_lines),
            ("Executive Summary", [tl_dr]),
            ("Discussion Points", key_points),
            ("Agenda / Structure", outline_lines),
            ("Action Items", action_items),
            ("Important Moments", timestamps),
            ("Full Transcript", [transcript_text]),
        ]

    return [
        ("Source", source_lines),
        ("TL;DR", [tl_dr]),
        ("Key Points", key_points),
        ("Outline", outline_lines),
        ("Action Items", action_items),
        ("Timestamps", timestamps),
        ("Transcript", [transcript_text]),
    ]


def _result_to_markdown(job: Job, template: str = "default") -> str:
    sections = _render_sections(job, template=template)
    lines = ["# AI Summary Result", ""]
    for title, items in sections:
        lines.append(f"## {title}")
        if title == "Transcript" or title == "Full Transcript":
            lines.append("```text")
            lines.append(items[0] if items else "")
            lines.append("```")
        else:
            for item in items:
                lines.append(f"- {item}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _result_to_docx_bytes(job: Job, template: str = "default") -> bytes:
    doc = Document()
    doc.add_heading("AI Summary Result", level=1)
    for title, items in _render_sections(job, template=template):
        doc.add_heading(title, level=2)
        if title in ("Transcript", "Full Transcript"):
            doc.add_paragraph(items[0] if items else "")
        else:
            for item in items:
                doc.add_paragraph(item, style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _result_to_pdf_bytes(job: Job, template: str = "default") -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x = 40
    y = height - 40

    def write_line(text: str, bold: bool = False):
        nonlocal y
        if y < 50:
            c.showPage()
            y = height - 40
        c.setFont("Helvetica-Bold" if bold else "Helvetica", 10)
        c.drawString(x, y, text[:140])
        y -= 14

    write_line("AI Summary Result", bold=True)
    write_line("")
    for title, items in _render_sections(job, template=template):
        write_line(title, bold=True)
        if title in ("Transcript", "Full Transcript"):
            text = items[0] if items else ""
            for paragraph in text.splitlines() or [""]:
                for chunk in textwrap.wrap(paragraph, width=120) or [""]:
                    write_line(chunk)
        else:
            for item in items:
                for chunk in textwrap.wrap(f"- {item}", width=120):
                    write_line(chunk)
        write_line("")

    c.save()
    return buf.getvalue()


@router.get("/config")
def get_config():
    """Public configuration (no auth needed)."""
    return {"max_upload_mb": settings.max_upload_mb}


@router.get("/{job_id}")
def get_job_status(
    job_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(404, "Job not found")

    return {
        "job_id": str(job.id),
        "status": job.status,
        "progress": job.progress,
        "source_type": job.source_type,
        "source_meta": job.source_meta,
        "error": job.error,
        "created_at": job.created_at.isoformat() if job.created_at else None,
    }


@router.get("")
def list_jobs(
    limit: int = 20,
    offset: int = 0,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    if limit < 1 or limit > 100:
        raise HTTPException(400, "limit must be between 1 and 100")
    if offset < 0:
        raise HTTPException(400, "offset must be >= 0")

    rows = (
        db.query(Job)
        .filter(Job.user_id == user.id)
        .order_by(Job.created_at.desc())
        .offset(offset)
        .limit(limit)
        .all()
    )
    items = [
        {
            "job_id": str(j.id),
            "status": j.status,
            "progress": j.progress,
            "source_type": j.source_type,
            "source_meta": j.source_meta,
            "error": j.error,
            "created_at": j.created_at.isoformat() if j.created_at else None,
        }
        for j in rows
    ]
    return {"items": items, "limit": limit, "offset": offset}


@router.get("/{job_id}/result")
def get_job_result(
    job_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(404, "Job not found")

    if job.status == "error":
        raise HTTPException(422, detail=job.error or {"message": "Processing failed"})
    if job.status != "done":
        raise HTTPException(409, "Job not done yet")

    return {
        "source": job.source_meta,
        "transcript": job.transcript,
        "summary": job.summary,
    }


@router.get("/{job_id}/result.md", response_class=PlainTextResponse)
def get_job_result_markdown(
    job_id: str,
    template: str = "default",
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(404, "Job not found")

    if job.status == "error":
        raise HTTPException(422, detail=job.error or {"message": "Processing failed"})
    if job.status != "done":
        raise HTTPException(409, "Job not done yet")

    filename = f"ai-summary-{job.id}.md"
    return PlainTextResponse(
        content=_result_to_markdown(job, template=template),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{job_id}/result.docx")
def get_job_result_docx(
    job_id: str,
    template: str = "default",
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(404, "Job not found")
    if job.status == "error":
        raise HTTPException(422, detail=job.error or {"message": "Processing failed"})
    if job.status != "done":
        raise HTTPException(409, "Job not done yet")

    data = _result_to_docx_bytes(job, template=template)
    filename = f"ai-summary-{job.id}.docx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{job_id}/result.pdf")
def get_job_result_pdf(
    job_id: str,
    template: str = "default",
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(404, "Job not found")
    if job.status == "error":
        raise HTTPException(422, detail=job.error or {"message": "Processing failed"})
    if job.status != "done":
        raise HTTPException(409, "Job not done yet")

    data = _result_to_pdf_bytes(job, template=template)
    filename = f"ai-summary-{job.id}.pdf"
    return Response(
        content=data,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/{job_id}/retry")
def retry_job(
    job_id: str,
    _: None = Depends(retry_limiter),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(404, "Job not found")
    if job.status != "error":
        raise HTTPException(409, "Only failed jobs can be retried")

    if job.source_type == "youtube":
        task = "app.workers.tasks.process_youtube"
    elif job.source_type == "upload":
        task = "app.workers.tasks.process_upload"
    else:
        raise HTTPException(400, f"Unsupported source_type: {job.source_type}")

    job.status = "queued"
    job.progress = 0
    job.error = None
    db.commit()

    enqueue_task(task, str(job.id))
    return {"job_id": str(job.id), "status": job.status}


@router.post("/{job_id}/cancel")
def cancel_job(
    job_id: str,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
    if not job:
        raise HTTPException(404, "Job not found")

    if job.status in ("done", "error", "cancelled"):
        raise HTTPException(409, f"Cannot cancel job in status: {job.status}")

    # If currently running, request cooperative cancellation in worker.
    if job.status == "running":
        job.status = "cancel_requested"
        db.commit()
        return {"job_id": str(job.id), "status": job.status}

    # If queued/scheduled, try to cancel in RQ immediately.
    q = get_queue()
    try:
        rq_job = RQJob.fetch(str(job.id), connection=q.connection)
        rq_job.cancel()
    except NoSuchJobError:
        pass

    job.status = "cancelled"
    job.progress = 0
    job.error = {"code": "cancelled", "message": "Cancelled by user", "retryable": False}
    db.commit()

    return {"job_id": str(job.id), "status": job.status}
